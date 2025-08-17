import streamlit as st
import tempfile
import os
import io
import time
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed

# Import availability flags and libraries
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import whisper
    WHISPER_AVAILABLE = True
except ImportError:
    WHISPER_AVAILABLE = False

try:
    from moviepy.editor import VideoFileClip
    MOVIEPY_AVAILABLE = True
except ImportError:
    MOVIEPY_AVAILABLE = False

try:
    from langdetect import detect
    LANGDETECT_AVAILABLE = True
except ImportError:
    LANGDETECT_AVAILABLE = False

from summarization import generate_summary

# File type mappings
SUPPORTED_FILE_TYPES = {
    'text': ['.txt', '.md', '.rtf'],
    'pdf': ['.pdf'],
    'word': ['.docx', '.doc'],
    'audio': ['.mp3', '.wav', '.m4a', '.ogg', '.flac'],
    'video': ['.mp4', '.avi', '.mov', '.mkv', '.wmv', '.flv']
}

def determine_file_type(file_extension):
    """
    Determine file type based on extension
    """
    file_ext = file_extension.lower()
    
    for file_type, extensions in SUPPORTED_FILE_TYPES.items():
        if file_ext in extensions:
            return file_type
    
    return 'text'  # Default fallback

def detect_language_safe(text):
    """
    Safely detect language with fallback
    """
    if not LANGDETECT_AVAILABLE or not text.strip():
        return "en"
    
    try:
        detected = detect(text)
        # Map to supported language codes
        language_mapping = {
            'en': 'en', 'es': 'es', 'fr': 'fr', 'de': 'de', 'it': 'it',
            'pt': 'pt', 'nl': 'nl', 'ru': 'ru', 'zh-cn': 'zh-cn', 
            'zh': 'zh-cn', 'ja': 'ja', 'ko': 'ko', 'ar': 'ar', 
            'hi': 'hi', 'tr': 'tr', 'pl': 'pl', 'sv': 'sv', 'no': 'no'
        }
        return language_mapping.get(detected, "en")
    except Exception:
        return "en"

def extract_text_from_pdf(file_bytes):
    """
    Extract text from PDF files using pdfplumber
    """
    if not PDF_AVAILABLE:
        raise Exception("PDF processing not available. Install with: pip install pdfplumber")
    
    text_content = []
    
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            total_pages = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_content.append(f"--- Page {page_num} ---\n{page_text}")
                
                # Show progress for large PDFs
                if total_pages > 10 and page_num % 5 == 0:
                    st.write(f"📄 Processing page {page_num}/{total_pages}...")
        
        combined_text = "\n\n".join(text_content)
        
        if not combined_text.strip():
            return "No readable text found in PDF. The file might contain only images or be corrupted."
        
        return combined_text
        
    except Exception as e:
        raise Exception(f"Error extracting PDF text: {str(e)}")

def extract_text_from_docx(file_bytes):
    """
    Extract text from Word documents
    """
    if not DOCX_AVAILABLE:
        raise Exception("Word document processing not available. Install with: pip install python-docx")
    
    try:
        doc = Document(io.BytesIO(file_bytes))
        
        # Extract paragraphs
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
        
        # Extract tables
        tables_text = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    table_data.append(" | ".join(row_data))
            
            if table_data:
                tables_text.append("\n".join(table_data))
        
        # Combine all text
        all_text = []
        if paragraphs:
            all_text.extend(paragraphs)
        if tables_text:
            all_text.append("\n--- Tables ---\n")
            all_text.extend(tables_text)
        
        combined_text = "\n\n".join(all_text)
        
        if not combined_text.strip():
            return "No readable text found in Word document."
        
        return combined_text
        
    except Exception as e:
        raise Exception(f"Error extracting Word document text: {str(e)}")

def extract_audio_from_video(video_bytes, output_path):
    """
    Extract audio from video files
    """
    if not MOVIEPY_AVAILABLE:
        raise Exception("Video processing not available. Install with: pip install moviepy")
    
    temp_video_path = None
    
    try:
        # Create temporary video file
        with tempfile.NamedTemporaryFile(suffix='.mp4', delete=False) as temp_video:
            temp_video.write(video_bytes)
            temp_video.flush()
            temp_video_path = temp_video.name
        
        # Extract audio
        video = VideoFileClip(temp_video_path)
        audio = video.audio
        
        if audio is None:
            raise Exception("No audio track found in video")
        
        audio.write_audiofile(output_path, verbose=False, logger=None)
        
        # Cleanup
        audio.close()
        video.close()
        
        return True
        
    except Exception as e:
        st.error(f"Video processing error: {str(e)}")
        return False
        
    finally:
        # Clean up temporary video file
        if temp_video_path and os.path.exists(temp_video_path):
            try:
                os.unlink(temp_video_path)
            except:
                pass

def transcribe_audio(audio_bytes, models):
    """
    Transcribe audio using Whisper
    """
    if not WHISPER_AVAILABLE or 'whisper' not in models:
        raise Exception("Audio transcription not available. Install with: pip install openai-whisper")
    
    temp_audio_path = None
    
    try:
        # Create temporary audio file
        with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_audio:
            temp_audio.write(audio_bytes)
            temp_audio.flush()
            temp_audio_path = temp_audio.name
        
        # Transcribe with Whisper
        with st.spinner("🎵 Transcribing audio..."):
            result = models['whisper'].transcribe(temp_audio_path)
        
        transcript = result['text']
        detected_language = result.get('language', 'en')
        
        if not transcript.strip():
            return "No speech detected in audio file.", 'en'
        
        return transcript, detected_language
        
    except Exception as e:
        raise Exception(f"Audio transcription error: {str(e)}")
        
    finally:
        # Clean up temporary audio file
        if temp_audio_path and os.path.exists(temp_audio_path):
            try:
                os.unlink(temp_audio_path)
            except:
                pass

def process_single_file(file_info, models, settings):
    """
    Process a single file and return results
    """
    start_time = time.time()
    
    result = {
        'filename': file_info['name'],
        'file_type': file_info['type'],
        'status': 'processing',
        'error': None,
        'transcript': '',
        'summary': '',
        'language': 'en',
        'processing_time': 0,
        'word_count': 0,
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }
    
    try:
        # Extract text based on file type
        if file_info['type'] == 'text':
            text_content = file_info['content']
            if isinstance(text_content, bytes):
                text_content = text_content.decode('utf-8', errors='ignore')
        
        elif file_info['type'] == 'pdf':
            if not PDF_AVAILABLE:
                raise Exception("PDF processing not available. Install pdfplumber.")
            text_content = extract_text_from_pdf(file_info['content'])
        
        elif file_info['type'] == 'word':
            if not DOCX_AVAILABLE:
                raise Exception("Word document processing not available. Install python-docx.")
            text_content = extract_text_from_docx(file_info['content'])
        
        elif file_info['type'] == 'audio':
            if not WHISPER_AVAILABLE:
                raise Exception("Audio processing not available. Install openai-whisper.")
            text_content, detected_lang = transcribe_audio(file_info['content'], models)
            result['language'] = detected_lang
        
        elif file_info['type'] == 'video':
            if not MOVIEPY_AVAILABLE or not WHISPER_AVAILABLE:
                raise Exception("Video processing not available. Install moviepy and openai-whisper.")
            
            with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_audio:
                if extract_audio_from_video(file_info['content'], temp_audio.name):
                    with open(temp_audio.name, 'rb') as audio_file:
                        audio_bytes = audio_file.read()
                    text_content, detected_lang = transcribe_audio(audio_bytes, models)
                    result['language'] = detected_lang
                    os.unlink(temp_audio.name)
                else:
                    raise Exception("Failed to extract audio from video")
        
        else:
            raise Exception(f"Unsupported file type: {file_info['type']}")
        
        # Store transcript and basic info
        result['transcript'] = text_content
        result['word_count'] = len(text_content.split()) if text_content else 0
        
        # Detect language for text-based files
        if file_info['type'] in ['text', 'pdf', 'word'] and settings.get('enable_language_detection', True):
            result['language'] = detect_language_safe(text_content)
        
        # Generate summary
        if text_content and text_content.strip():
            summary = generate_summary(
                text_content,
                settings.get('model', 'BART'),
                models,
                settings
            )
            result['summary'] = summary
        else:
            result['summary'] = "No content found to summarize."
        
        result['status'] = 'completed'
        
    except Exception as e:
        result['status'] = 'error'
        result['error'] = str(e)
    
    result['processing_time'] = time.time() - start_time
    return result

def process_batch_files(uploaded_files, models, settings):
    """
    Process multiple files concurrently
    """
    # Prepare file information
    file_infos = []
    for file in uploaded_files:
        file_ext = os.path.splitext(file.name)[1].lower()
        file_type = determine_file_type(file_ext)
        
        file_infos.append({
            'name': file.name,
            'type': file_type,
            'content': file.read()
        })
    
    # Create progress containers
    progress_bar = st.progress(0)
    status_container = st.empty()
    results_container = st.container()
    
    # Concurrent processing
    max_workers = min(len(file_infos), settings.get('batch_size', 4))
    results = []
    
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        # Submit all tasks
        future_to_file = {
            executor.submit(process_single_file, file_info, models, settings): file_info 
            for file_info in file_infos
        }
        
        completed = 0
        for future in as_completed(future_to_file):
            file_info = future_to_file[future]
            completed += 1
            
            try:
                result = future.result()
                results.append(result)
                
                # Update progress
                progress = completed / len(file_infos)
                progress_bar.progress(progress)
                
                # Status update
                if result['status'] == 'completed':
                    status_container.success(
                        f"✅ {file_info['name']} completed in {result['processing_time']:.1f}s"
                    )
                else:
                    status_container.error(f"❌ {file_info['name']}: {result['error']}")
                
            except Exception as e:
                st.error(f"Processing error for {file_info['name']}: {e}")
    
    # Celebration for successful completion
    successful = len([r for r in results if r['status'] == 'completed'])
    if successful > 0:
        st.balloons()
        st.success(f"🎉 Batch processing completed! {successful}/{len(results)} files processed successfully")
    
    return results
